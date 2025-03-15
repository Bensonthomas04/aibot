import os
from dotenv import load_dotenv
from telegram import Update, InputFile
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
import nest_asyncio  # For environments where an event loop is already running
import random  # For introducing randomness in text rewriting

# Apply nest_asyncio to allow re-entrant event loops (for Jupyter/IDEs)
nest_asyncio.apply()

# Load environment variables
load_dotenv()

# Get the Telegram Bot Token from .env
TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')

# Validate environment variables
if not TELEGRAM_BOT_TOKEN:
    raise ValueError("Please ensure TELEGRAM_BOT_TOKEN is set in the .env file.")

# Command: /start
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Welcome! Send me a .docx file, and I'll rewrite the text while keeping the formatting intact."
    )

# Handle incoming .docx files
async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    # Check if the file is a .docx
    if not update.message.document.file_name.endswith('.docx'):
        await update.message.reply_text("Please send a .docx file.")
        return

    # Download the file
    file = await update.message.document.get_file()
    file_path = f"temp_{update.message.document.file_name}"
    await file.download_to_drive(file_path)

    # Process the file
    processed_file_path = None
    try:
        processed_file_path = process_docx(file_path)
        
        # Send the processed file back to the user
        with open(processed_file_path, 'rb') as f:
            await update.message.reply_document(document=InputFile(f))
        
        # Notify the user
        await update.message.reply_text("Your file has been processed and plagiarism has been reduced.")

    except Exception as e:
        await update.message.reply_text(f"An error occurred: {e}")

    finally:
        # Clean up temporary files
        if os.path.exists(file_path):
            os.remove(file_path)
        if processed_file_path and os.path.exists(processed_file_path):
            os.remove(processed_file_path)

# Process the .docx file
def process_docx(file_path: str) -> str:
    # Read the .docx file
    doc = Document(file_path)

    # Create a new document to store the rewritten content
    new_doc = Document()

    # Iterate through each paragraph in the original document
    for para in doc.paragraphs:
        # Preserve paragraph style (e.g., headings, bullet points)
        new_para = new_doc.add_paragraph(style=para.style)

        # Rewrite the text while preserving formatting
        for run in para.runs:
            rewritten_text = rewrite_text(run.text)  # Rewrite the text
            new_run = new_para.add_run(rewritten_text)

            # Preserve formatting (bold, italic, underline, font size, font type)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size

    # Save the new document
    new_file_path = f"processed_{os.path.basename(file_path)}"
    new_doc.save(new_file_path)

    return new_file_path

# Advanced rule-based text rewriting
def rewrite_text(text: str) -> str:
    # Split the text into sentences
    sentences = text.split(". ")

    # Rewrite each sentence
    rewritten_sentences = []
    for sentence in sentences:
        if sentence.strip():  # Skip empty sentences
            rewritten_sentence = rewrite_sentence(sentence)
            rewritten_sentences.append(rewritten_sentence)

    # Join the rewritten sentences
    return ". ".join(rewritten_sentences)

# Rewrite a single sentence
def rewrite_sentence(sentence: str) -> str:
    # Split the sentence into words
    words = sentence.split()

    # Apply rewriting rules
    rewritten_words = []
    for word in words:
        # Randomly replace words with synonyms
        if random.random() < 0.3:  # 30% chance to replace a word
            synonym = get_synonym(word)
            rewritten_words.append(synonym)
        else:
            rewritten_words.append(word)

    # Randomly shuffle the order of words (optional)
    if random.random() < 0.2:  # 20% chance to shuffle words
        random.shuffle(rewritten_words)

    # Join the words into a sentence
    return " ".join(rewritten_words)

# Get a synonym for a word (simple example)
def get_synonym(word: str) -> str:
    synonym_map = {
        "student": "learner",
        "university": "college",
        "education": "learning",
        "knowledge": "understanding",
        "research": "investigation",
        "study": "analysis",
        "assignment": "task",
        "project": "undertaking",
        "professor": "lecturer",
        "lecture": "class",
        "campus": "grounds",
        "degree": "qualification",
        "course": "program",
        "textbook": "manual",
        "library": "archive",
        "exam": "test",
        "grade": "mark",
        "scholarship": "grant",
        "tuition": "fee",
        "dormitory": "hostel",
        "roommate": "housemate",
        "cafeteria": "dining hall",
        "graduation": "commencement",
        "diploma": "certificate",
        "thesis": "dissertation",
        "internship": "apprenticeship",
        "semester": "term",
        "curriculum": "syllabus",
        "faculty": "staff",
        "department": "division",
        "major": "specialization",
        "minor": "secondary focus",
        "lecturer": "instructor",
        "tutor": "mentor",
        "peer": "colleague",
        "assignment": "homework",
        "deadline": "due date",
        "plagiarism": "academic dishonesty",
        "citation": "reference",
        "bibliography": "reference list",
        "abstract": "summary",
        "hypothesis": "assumption",
        "methodology": "approach",
        "data": "information",
        "analysis": "evaluation",
        "conclusion": "finding",
        "argument": "claim",
        "evidence": "proof",
        "theory": "concept",
        "principle": "rule",
        "phenomenon": "occurrence",
        "variable": "factor",
        "experiment": "trial",
        "observation": "monitoring",
        "result": "outcome",
        "discussion": "debate",
        "recommendation": "suggestion",
        "limitation": "constraint",
        "implication": "consequence",
        "framework": "structure",
        "model": "representation",
        "paradigm": "pattern",
        "perspective": "viewpoint",
        "context": "background",
        "scope": "range",
        "objective": "goal",
        "strategy": "plan",
        "technique": "method",
        "tool": "instrument",
        "resource": "asset",
        "contribution": "input",
        "impact": "effect",
        "trend": "pattern",
        "issue": "problem",
        "challenge": "difficulty",
        "solution": "answer",
        "innovation": "invention",
        "creativity": "originality",
        "collaboration": "cooperation",
        "communication": "interaction",
        "presentation": "delivery",
        "publication": "release",
        "journal": "periodical",
        "article": "paper",
        "author": "writer",
        "editor": "reviewer",
        "review": "evaluation",
        "feedback": "response",
        "revision": "amendment",
        "draft": "version",
        "final": "completed",
        "submission": "entry",
        "approval": "acceptance",
        "rejection": "denial",
        "criticism": "critique",
        "improvement": "enhancement",
        "progress": "advancement",
        "achievement": "accomplishment",
        "success": "triumph",
        "failure": "defeat",
        "motivation": "drive",
        "inspiration": "stimulus",
        "dedication": "commitment",
        "effort": "endeavor",
        "hard work": "diligence",
        "persistence": "determination",
        "resilience": "toughness",
        "adaptability": "flexibility",
        "initiative": "enterprise",
        "leadership": "guidance",
        "teamwork": "collaboration",
        "ethics": "morality",
        "integrity": "honesty",
        "responsibility": "accountability",
        "discipline": "self-control",
        "time management": "scheduling",
        "organization": "arrangement",
        "priority": "importance",
        "efficiency": "productivity",
        "effectiveness": "competence",
        "quality": "standard",
        "excellence": "superiority",
        "performance": "execution",
        "assessment": "appraisal",
        "evaluation": "judgment",
        "measurement": "quantification",
        "benchmark": "standard",
        "indicator": "signal",
        "criterion": "requirement",
        "standard": "norm",
        "expectation": "anticipation",
        "satisfaction": "fulfillment",
        "dissatisfaction": "discontent",
        "frustration": "disappointment",
        "stress": "pressure",
        "anxiety": "worry",
        "pressure": "strain",
        "burnout": "exhaustion",
        "balance": "equilibrium",
        "well-being": "health",
        "mental health": "psychological state",
        "physical health": "bodily condition",
        "nutrition": "diet",
        "exercise": "workout",
        "sleep": "rest",
        "relaxation": "unwinding",
        "hobby": "pastime",
        "interest": "curiosity",
        "passion": "enthusiasm",
        "goal": "objective",
        "dream": "aspiration",
        "ambition": "desire",
        "career": "profession",
        "job": "employment",
        "work": "labor",
        "employer": "boss",
        "employee": "worker",
        "colleague": "coworker",
        "supervisor": "manager",
        "mentor": "guide",
        "role model": "example",
        "network": "connections",
        "opportunity": "chance",
        "challenge": "obstacle",
        "competition": "rivalry",
        "market": "industry",
        "economy": "financial system",
        "globalization": "internationalization",
        "technology": "innovation",
        "digital": "electronic",
        "internet": "web",
        "social media": "online platforms",
        "communication": "interaction",
        "information": "data",
        "knowledge": "understanding",
        "learning": "education",
        "skill": "ability",
        "experience": "practice",
        "expertise": "proficiency",
        "competence": "capability",
        "qualification": "credential",
        "certification": "accreditation",
        "license": "permit",
        "training": "instruction",
        "development": "growth",
        "improvement": "enhancement",
        "progress": "advancement",
        "success": "achievement",
        "failure": "setback",
        "mistake": "error",
        "lesson": "teaching",
        "feedback": "response",
        "criticism": "evaluation",
        "praise": "compliment",
        "recognition": "acknowledgment",
        "reward": "prize",
        "punishment": "penalty",
        "discipline": "control",
        "motivation": "drive",
        "inspiration": "stimulus",
        "creativity": "originality",
        "innovation": "invention",
        "problem": "issue",
        "solution": "answer",
        "decision": "choice",
        "strategy": "plan",
        "plan": "scheme",
        "goal": "objective",
        "objective": "aim",
        "purpose": "intention",
        "mission": "goal",
        "vision": "dream",
        "value": "principle",
        "belief": "conviction",
        "attitude": "mindset",
        "behavior": "conduct",
        "habit": "routine",
        "culture": "tradition",
        "diversity": "variety",
        "inclusion": "integration",
        "equality": "fairness",
        "justice": "fairness",
        "right": "entitlement",
        "responsibility": "duty",
        "ethics": "morality",
        "integrity": "honesty",
        "trust": "confidence",
        "respect": "esteem",
        "honor": "dignity",
        "reputation": "standing",
        "image": "perception",
        "identity": "self",
        "personality": "character",
        "emotion": "feeling",
        "mood": "temperament",
        "happiness": "joy",
        "sadness": "sorrow",
        "anger": "fury",
        "fear": "anxiety",
        "love": "affection",
        "hate": "dislike",
        "friendship": "companionship",
        "relationship": "connection",
        "family": "household",
        "parent": "guardian",
        "child": "offspring",
        "sibling": "brother/sister",
        "partner": "companion",
        "marriage": "union",
        "divorce": "separation",
        "community": "society",
        "neighbor": "local",
        "citizen": "resident",
        "government": "administration",
        "politics": "governance",
        "policy": "regulation",
        "law": "rule",
        "justice": "fairness",
        "crime": "offense",
        "punishment": "penalty",
        "freedom": "liberty",
        "right": "entitlement",
        "responsibility": "duty",
        "duty": "obligation",
        "service": "assistance",
        "volunteer": "helper",
        "charity": "philanthropy",
        "donation": "contribution",
        "support": "assistance",
        "help": "aid",
        "care": "attention",
        "health": "well-being",
        "medicine": "treatment",
        "doctor": "physician",
        "patient": "sufferer",
        "hospital": "clinic",
        "disease": "illness",
        "symptom": "indication",
        "treatment": "therapy",
        "recovery": "healing",
        "prevention": "avoidance",
        "vaccine": "immunization",
        "epidemic": "outbreak",
        "pandemic": "global outbreak",
        "environment": "surroundings",
        "nature": "wildlife",
        "pollution": "contamination",
        "climate": "weather",
        "change": "transformation",
        "global warming": "climate change",
        "sustainability": "endurance",
        "conservation": "preservation",
        "energy": "power",
        "resource": "asset",
        "waste": "garbage",
        "recycling": "reuse",
        "technology": "innovation",
        "science": "knowledge",
        "research": "investigation",
        "discovery": "finding",
        "invention": "creation",
        "experiment": "trial",
        "observation": "monitoring",
        "data": "information",
        "analysis": "evaluation",
        "conclusion": "result",
        "theory": "hypothesis",
        "principle": "rule",
        "law": "regulation",
        "fact": "truth",
        "evidence": "proof",
        "argument": "claim",
        "debate": "discussion",
        "opinion": "view",
        "belief": "conviction",
        "truth": "reality",
        "lie": "falsehood",
        "honesty": "integrity",
        "trust": "confidence",
        "doubt": "skepticism",
        "certainty": "confidence",
        "uncertainty": "doubt",
        "risk": "danger",
        "safety": "security",
        "danger": "hazard",
        "threat": "risk",
        "protection": "defense",
        "security": "safety",
        "fear": "anxiety",
        "courage": "bravery",
        "hope": "optimism",
        "despair": "hopelessness",
        "faith": "belief",
        "religion": "spirituality",
        "spirituality": "faith",
        "prayer": "meditation",
        "worship": "devotion",
        "god": "deity",
        "angel": "messenger",
        "devil": "demon",
        "heaven": "paradise",
        "hell": "underworld",
        "soul": "spirit",
        "life": "existence",
        "death": "demise",
        "birth": "beginning",
        "rebirth": "reincarnation",
        "eternity": "infinity",
        "time": "duration",
        "past": "history",
        "present": "now",
        "future": "tomorrow",
        "age": "era",
        "generation": "epoch",
        "century": "hundred years",
        "decade": "ten years",
        "year": "twelve months",
        "month": "four weeks",
        "week": "seven days",
        "day": "twenty-four hours",
        "hour": "sixty minutes",
        "minute": "sixty seconds",
        "second": "moment",
        "moment": "instant",
        "history": "past",
        "tradition": "custom",
        "culture": "heritage",
        "art": "creativity",
        "music": "melody",
        "dance": "movement",
        "theater": "drama",
        "film": "movie",
        "literature": "writing",
        "poetry": "verse",
        "story": "tale",
        "novel": "book",
        "author": "writer",
        "reader": "audience",
        "language": "tongue",
        "word": "term",
        "sentence": "phrase",
        "paragraph": "section",
        "chapter": "part",
        "book": "volume",
        "library": "archive",
        "knowledge": "wisdom",
        "wisdom": "insight",
        "intelligence": "smarts",
        "smart": "clever",
        "stupid": "foolish",
        "genius": "prodigy",
        "idiot": "fool",
        "expert": "specialist",
        "amateur": "beginner",
        "professional": "expert",
        "work": "labor",
        "job": "occupation",
        "career": "profession",
        "business": "enterprise",
        "company": "firm",
        "organization": "institution",
        "team": "group",
        "leader": "manager",
        "manager": "supervisor",
        "boss": "employer",
        "employee": "worker",
        "colleague": "coworker",
        "partner": "associate",
        "client": "customer",
        "customer": "buyer",
        "consumer": "user",
        "product": "item",
        "service": "assistance",
        "price": "cost",
        "cost": "expense",
        "value": "worth",
        "profit": "gain",
        "loss": "deficit",
        "money": "currency",
        "wealth": "riches",
        "poverty": "destitution",
        "rich": "wealthy",
        "poor": "needy",
        "economy": "financial system",
        "market": "industry",
        "trade": "commerce",
        "export": "shipment",
        "import": "purchase",
        "investment": "funding",
        "stock": "share",
        "bond": "security",
        "bank": "financial institution",
        "loan": "credit",
        "debt": "liability",
        "interest": "return",
        "tax": "levy",
        "income": "earnings",
        "salary": "wage",
        "wage": "pay",
        "payment": "remittance",
        "bill": "invoice",
        "expense": "cost",
        "budget": "plan",
        "saving": "reserve",
        "spending": "expenditure",
        "wealth": "fortune",
        "poverty": "hardship",
        "inequality": "disparity",
        "justice": "fairness",
        "injustice": "unfairness",
        "corruption": "dishonesty",
        "scandal": "controversy",
        "crime": "offense",
        "law": "regulation",
        "police": "law enforcement",
        "court": "tribunal",
        "judge": "magistrate",
        "lawyer": "attorney",
        "trial": "hearing",
        "verdict": "decision",
        "guilty": "culpable",
        "innocent": "blameless",
        "punishment": "penalty",
        "prison": "jail",
        "freedom": "liberty",
        "slavery": "bondage",
        "war": "conflict",
        "peace": "harmony",
        "violence": "aggression",
        "terrorism": "extremism",
        "attack": "assault",
        "defense": "protection",
        "soldier": "warrior",
        "army": "military",
        "weapon": "armament",
        "bomb": "explosive",
        "gun": "firearm",
        "knife": "blade",
        "fight": "battle",
        "victory": "triumph",
        "defeat": "loss",
        "enemy": "foe",
        "ally": "partner",
        "friend": "companion",
        "stranger": "unknown",
        "neighbor": "local",
        "community": "society",
        "nation": "country",
        "state": "province",
        "city": "town",
        "village": "hamlet",
        "home": "residence",
        "house": "dwelling",
        "apartment": "flat",
        "room": "chamber",
        "kitchen": "cooking area",
        "bathroom": "washroom",
        "bedroom": "sleeping area",
        "living room": "sitting room",
        "furniture": "appliance",
        "table": "desk",
        "chair": "seat",
        "bed": "cot",
        "sofa": "couch",
        "lamp": "light",
        "window": "opening",
        "door": "entrance",
        "wall": "barrier",
        "floor": "ground",
        "ceiling": "roof",
        "garden": "yard",
        "tree": "plant",
        "flower": "blossom",
        "grass": "lawn",
        "animal": "creature",
        "dog": "canine",
        "cat": "feline",
        "bird": "avian",
        "fish": "aquatic animal",
        "horse": "equine",
        "cow": "bovine",
        "sheep": "ovine",
        "pig": "swine",
        "chicken": "poultry",
        "egg": "ovum",
        "milk": "dairy",
        "meat": "flesh",
        "vegetable": "plant",
        "fruit": "produce",
        "grain": "cereal",
        "bread": "loaf",
        "rice": "staple",
        "pasta": "noodles",
        "soup": "broth",
        "salad": "greens",
        "sandwich": "snack",
        "pizza": "pie",
        "burger": "sandwich",
        "fries": "chips",
        "dessert": "sweet",
        "cake": "pastry",
        "cookie": "biscuit",
        "chocolate": "candy",
        "ice cream": "frozen dessert",
        "drink": "beverage",
        "water": "H2O",
        "juice": "liquid",
        "soda": "pop",
        "coffee": "brew",
        "tea": "infusion",
        "alcohol": "liquor",
        "beer": "ale",
        "wine": "vino",
        "whiskey": "spirit",
        "vodka": "liquor",
        "restaurant": "eatery",
        "cafe": "coffee shop",
        "bar": "pub",
        "hotel": "inn",
        "motel": "lodging",
        "resort": "retreat",
        "vacation": "holiday",
        "travel": "journey",
        "trip": "excursion",
        "flight": "air travel",
        "airport": "terminal",
        "train": "rail",
        "station": "depot",
        "bus": "coach",
        "car": "automobile",
        "bike": "bicycle",
        "motorcycle": "motorbike",
        "truck": "lorry",
        "ship": "vessel",
        "boat": "craft",
        "plane": "aircraft",
    }
    return synonym_map.get(word.lower(), word)

# Main function to start the bot
async def main() -> None:
    # Create the Application
    application = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    # Add handlers
    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.Document.ALL, handle_document))

    # Start the bot
    await application.run_polling()

if __name__ == '__main__':
    import asyncio

    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("Bot stopped.")